from docx import Document as DocxDocument
from langchain_core.documents import Document

from src.pyxon.parsers.base import BaseParser


class PyxonDocxParser(BaseParser):
    SUPPORTED_EXTENSIONS: list[str] = [".doc", ".docx"]

    def __init__(self, file_path):
        super().__init__(file_path)

    def parse(self) -> Document:
        doc = DocxDocument(self._file_path)
        paragraphs = [p.text for p in doc.paragraphs]

        content = "\n".join(paragraphs)

        return Document(page_content=content, metadata={"source": str(self._file_path)})
